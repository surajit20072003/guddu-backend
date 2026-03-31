import io
import re

import docx
import pdfplumber


CHAPTER_PATTERNS = [
    re.compile(r"^\s*(chapter|chap|unit)\s*[-:\.]?\s*(\d+)\b", re.IGNORECASE),
    re.compile(r"^\s*(\d+)\s*[-:\.]\s*(.+)$"),
]
TOPIC_BULLET_PATTERN = re.compile(r"^\s*(?:[-*]|[0-9]+[\.\)]|[a-zA-Z][\.\)])\s+(.+)$")


def validate_upload(file_obj, max_size_bytes=10 * 1024 * 1024):
    if not file_obj:
        raise ValueError("file is required")

    filename = getattr(file_obj, "name", "") or ""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in {"pdf", "doc", "docx"}:
        raise ValueError("Unsupported file type. Allowed: PDF, DOC, DOCX")

    size = getattr(file_obj, "size", None)
    if size is not None and size > max_size_bytes:
        raise ValueError("File too large. Maximum size is 10MB")

    return ext


def extract_text_from_uploaded_file(file_obj):
    ext = validate_upload(file_obj)
    file_obj.seek(0)
    raw = file_obj.read()
    if not raw:
        raise ValueError("Uploaded file is empty")

    text = ""
    if ext == "pdf":
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
    elif ext == "docx":
        document = docx.Document(io.BytesIO(raw))
        for para in document.paragraphs:
            text += para.text + "\n"
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
    else:  # .doc basic fallback
        text = raw.decode("utf-8", errors="ignore")

    text = text.strip()
    if not text:
        raise ValueError("Could not extract readable text from the file")
    return text


def parse_syllabus_text(text):
    """
    Parse raw text into a structure:
    [{"chapter_number": 1, "title": "...", "topics": ["...", "..."]}, ...]
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    chapters = []
    current = None
    chapter_counter = 0

    for line in lines:
        chapter_match = None
        chapter_num = None
        chapter_title = None

        for pattern in CHAPTER_PATTERNS:
            chapter_match = pattern.match(line)
            if chapter_match:
                groups = chapter_match.groups()
                if len(groups) >= 2 and str(groups[1]).isdigit():
                    chapter_num = int(groups[1])
                    chapter_title = line
                elif len(groups) >= 2 and str(groups[0]).isdigit():
                    chapter_num = int(groups[0])
                    chapter_title = groups[1].strip() or line
                else:
                    chapter_title = line
                break

        if chapter_match:
            chapter_counter += 1
            current = {
                "chapter_number": chapter_num or chapter_counter,
                "title": chapter_title or line,
                "topics": [],
            }
            chapters.append(current)
            continue

        topic_match = TOPIC_BULLET_PATTERN.match(line)
        topic_text = topic_match.group(1).strip() if topic_match else line

        if current is None:
            chapter_counter += 1
            current = {"chapter_number": chapter_counter, "title": f"Chapter {chapter_counter}", "topics": []}
            chapters.append(current)

        if topic_text and topic_text not in current["topics"]:
            current["topics"].append(topic_text)

    if not chapters:
        raise ValueError("No chapters found in file content")
    return chapters
