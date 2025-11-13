# api/parsers.py

import pdfplumber # <-- The new, smart PDF module
import openpyxl
import docx
import re # <-- Import regular expressions

# --- FUNCTION 1: Read text from files (SMARTER PDF VERSION) ---
def get_keywords_from_file(file_path):
    text = ""
    try:
        if file_path.endswith('.pdf'):
            # Use pdfplumber to read the PDF
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # extract_text() is smarter and respects layout
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        
        elif file_path.endswith('.xlsx'):
            wb = openpyxl.load_workbook(file_path, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    row_text = ""
                    for cell in row:
                        if cell.value:
                            row_text += str(cell.value) + " "
                    text += row_text.strip() + "\n"
                            
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            # 1. Read from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n" # Use newline
            # 2. Read from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " "
                    text += row_text.strip() + "\n"
                        
    except Exception as e:
        print(f"Error parsing file {file_path}: {e}")
        
    print(f"Extracted {len(text)} chars from file.")
    return text

# --- FUNCTION 2: Clean text and extract keywords (YOUR LOGIC) ---
def extract_keywords_from_text(text):
    """
    Extracts keywords based on commas and multiple spaces.
    """
    if not text:
        return []

    potential_tags = []
    
    # This pattern will find things like "a)", "b)", "1.", "i)", "Sl. No."
    list_item_pattern = re.compile(r'^\s*(\w[\.\)]|\d+[\.\)]|Sl\.?\s*No\.?)')
    
    # 1. Normalize text: Treat all newlines as regular spaces
    text = text.replace('\n', ' ')

    # 2. Find delimiters: Use commas and 2+ spaces as separators
    # We replace them with a unique pipe "|" character
    text = text.replace(',', '|')
    text = re.sub(r'\s{2,}', '|', text) # \s{2,} = 2 or more spaces

    # 3. Now, split the text by our unique separator
    phrases = text.split('|')
    
    for phrase in phrases:
        # 4. Clean it
        clean_phrase = phrase.strip()

        # --- Filter Rules ---
        
        # Rule 1: Length
        if not (4 < len(clean_phrase) < 100):
            continue
            
        # Rule 2: Noise
        if 'www.' in clean_phrase or 'http' in clean_phrase or 'Topics Covered' in clean_phrase:
            continue
            
        # Rule 3: List items
        if list_item_pattern.match(clean_phrase):
            continue
            
        # Rule 4: Remove "for lkg" etc. (simple version)
        if clean_phrase.lower().startswith('for '):
            continue

        # If it passes all rules, it's a good tag
        potential_tags.append(clean_phrase)

    # 5. Remove duplicates
    final_tags = list(set(potential_tags))
    print(f"Found tags: {final_tags}") # For debugging
    return final_tags